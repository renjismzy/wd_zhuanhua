#!/usr/bin/env python3
"""
Document Converter Module

Provides asynchronous document conversion capabilities between various formats.
"""

import asyncio
import logging
import os
import tempfile
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

import aiofiles
import markdown
from docx import Document
from PyPDF2 import PdfReader
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch

logger = logging.getLogger(__name__)


class DocumentConverter:
    """文档转换器类"""
    
    def __init__(self):
        self.supported_formats = {
            'txt': {'read': True, 'write': True},
            'md': {'read': True, 'write': True},
            'html': {'read': True, 'write': True},
            'pdf': {'read': True, 'write': True},
            'docx': {'read': True, 'write': True},
            'doc': {'read': True, 'write': False}  # 只读
        }
        self.jobs = {}  # 存储转换任务状态
        
    def get_supported_formats(self) -> List[str]:
        """获取支持的格式列表"""
        return list(self.supported_formats.keys())
    
    async def convert_async(
        self,
        source_format: str,
        target_format: str,
        content: Optional[str] = None,
        file_path: Optional[str] = None
    ) -> str:
        """异步文档转换"""
        job_id = str(uuid.uuid4())
        
        # 初始化任务状态
        self.jobs[job_id] = {
            'status': 'pending',
            'progress': 0,
            'source_format': source_format,
            'target_format': target_format,
            'created_at': datetime.now(),
            'error': None,
            'result': None
        }
        
        # 启动异步转换任务
        asyncio.create_task(self._convert_task(job_id, source_format, target_format, content, file_path))
        
        return job_id
    
    async def _convert_task(
        self,
        job_id: str,
        source_format: str,
        target_format: str,
        content: Optional[str] = None,
        file_path: Optional[str] = None
    ):
        """执行转换任务"""
        try:
            self.jobs[job_id]['status'] = 'running'
            self.jobs[job_id]['progress'] = 10
            
            # 验证格式支持
            if source_format not in self.supported_formats:
                raise ValueError(f"Unsupported source format: {source_format}")
            if target_format not in self.supported_formats:
                raise ValueError(f"Unsupported target format: {target_format}")
            if not self.supported_formats[target_format]['write']:
                raise ValueError(f"Cannot write to format: {target_format}")
            
            self.jobs[job_id]['progress'] = 30
            
            # 读取源文档
            if file_path:
                source_content = await self._read_file(file_path, source_format)
            else:
                source_content = content
            
            if not source_content:
                raise ValueError("No content to convert")
            
            self.jobs[job_id]['progress'] = 60
            
            # 执行转换
            result = await self._convert_content(source_content, source_format, target_format)
            
            self.jobs[job_id]['progress'] = 90
            
            # 保存结果
            if target_format in ['pdf', 'docx']:
                # 二进制格式需要保存到文件
                output_path = await self._save_binary_result(result, target_format)
                self.jobs[job_id]['result'] = {
                    'success': True,
                    'file_path': output_path,
                    'content_type': self._get_content_type(target_format)
                }
            else:
                # 文本格式直接返回内容
                self.jobs[job_id]['result'] = {
                    'success': True,
                    'content': result,
                    'content_type': self._get_content_type(target_format)
                }
            
            self.jobs[job_id]['status'] = 'completed'
            self.jobs[job_id]['progress'] = 100
            
        except Exception as e:
            logger.error(f"Conversion error for job {job_id}: {e}")
            self.jobs[job_id]['status'] = 'failed'
            self.jobs[job_id]['error'] = str(e)
    
    async def _read_file(self, file_path: str, format: str) -> str:
        """读取文件内容"""
        if format == 'txt' or format == 'md' or format == 'html':
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                return await f.read()
        elif format == 'pdf':
            return await self._read_pdf(file_path)
        elif format == 'docx':
            return await self._read_docx(file_path)
        else:
            raise ValueError(f"Cannot read format: {format}")
    
    async def _read_pdf(self, file_path: str) -> str:
        """读取PDF文件"""
        def _read_pdf_sync():
            with open(file_path, 'rb') as file:
                reader = PdfReader(file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                return text
        
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, _read_pdf_sync)
    
    async def _read_docx(self, file_path: str) -> str:
        """读取DOCX文件"""
        def _read_docx_sync():
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, _read_docx_sync)
    
    async def _convert_content(self, content: str, source_format: str, target_format: str) -> Union[str, bytes]:
        """转换文档内容"""
        if source_format == target_format:
            return content
        
        # 转换为中间格式（纯文本或HTML）
        intermediate_content = await self._to_intermediate(content, source_format)
        
        # 从中间格式转换为目标格式
        return await self._from_intermediate(intermediate_content, target_format)
    
    async def _to_intermediate(self, content: str, source_format: str) -> str:
        """转换为中间格式（HTML）"""
        if source_format == 'txt':
            # 纯文本转HTML，保留换行
            return content.replace('\n', '<br>\n')
        elif source_format == 'md':
            # Markdown转HTML
            return markdown.markdown(content)
        elif source_format == 'html':
            return content
        elif source_format in ['pdf', 'docx']:
            # 已经是纯文本，转换为HTML
            return content.replace('\n', '<br>\n')
        else:
            return content
    
    async def _from_intermediate(self, content: str, target_format: str) -> Union[str, bytes]:
        """从中间格式转换为目标格式"""
        if target_format == 'txt':
            # HTML转纯文本
            import re
            # 简单的HTML标签移除
            text = re.sub(r'<br\s*/?>', '\n', content)
            text = re.sub(r'<[^>]+>', '', text)
            return text.strip()
        elif target_format == 'md':
            # HTML转Markdown（简化版）
            import re
            text = content
            text = re.sub(r'<h1[^>]*>(.*?)</h1>', r'# \1', text)
            text = re.sub(r'<h2[^>]*>(.*?)</h2>', r'## \1', text)
            text = re.sub(r'<h3[^>]*>(.*?)</h3>', r'### \1', text)
            text = re.sub(r'<strong[^>]*>(.*?)</strong>', r'**\1**', text)
            text = re.sub(r'<em[^>]*>(.*?)</em>', r'*\1*', text)
            text = re.sub(r'<br\s*/?>', '\n', text)
            text = re.sub(r'<[^>]+>', '', text)
            return text.strip()
        elif target_format == 'html':
            return content
        elif target_format == 'pdf':
            return await self._create_pdf(content)
        elif target_format == 'docx':
            return await self._create_docx(content)
        else:
            return content
    
    async def _create_pdf(self, html_content: str) -> bytes:
        """创建PDF文件"""
        def _create_pdf_sync():
            from io import BytesIO
            import re
            
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            # 简单的HTML解析和转换
            text = re.sub(r'<br\s*/?>', '\n', html_content)
            text = re.sub(r'<[^>]+>', '', text)
            
            for line in text.split('\n'):
                if line.strip():
                    p = Paragraph(line, styles['Normal'])
                    story.append(p)
                    story.append(Spacer(1, 0.2*inch))
            
            doc.build(story)
            buffer.seek(0)
            return buffer.getvalue()
        
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, _create_pdf_sync)
    
    async def _create_docx(self, html_content: str) -> bytes:
        """创建DOCX文件"""
        def _create_docx_sync():
            from io import BytesIO
            import re
            
            doc = Document()
            
            # 简单的HTML解析和转换
            text = re.sub(r'<br\s*/?>', '\n', html_content)
            text = re.sub(r'<[^>]+>', '', text)
            
            for line in text.split('\n'):
                if line.strip():
                    doc.add_paragraph(line)
            
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()
        
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, _create_docx_sync)
    
    async def _save_binary_result(self, data: bytes, format: str) -> str:
        """保存二进制结果到临时文件"""
        suffix = f'.{format}'
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp_file:
            tmp_file.write(data)
            return tmp_file.name
    
    def _get_content_type(self, format: str) -> str:
        """获取内容类型"""
        content_types = {
            'txt': 'text/plain',
            'md': 'text/markdown',
            'html': 'text/html',
            'pdf': 'application/pdf',
            'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        }
        return content_types.get(format, 'application/octet-stream')
    
    async def get_job_status(self, job_id: str) -> Dict[str, Any]:
        """获取任务状态"""
        if job_id not in self.jobs:
            return {'error': 'Job not found'}
        
        job = self.jobs[job_id].copy()
        if 'created_at' in job:
            job['created_at'] = job['created_at'].isoformat()
        
        return job
    
    async def wait_for_completion(self, job_id: str, timeout: int = 300) -> Dict[str, Any]:
        """等待任务完成"""
        start_time = asyncio.get_event_loop().time()
        
        while True:
            if job_id not in self.jobs:
                return {'success': False, 'error': 'Job not found'}
            
            job = self.jobs[job_id]
            
            if job['status'] == 'completed':
                return job['result']
            elif job['status'] == 'failed':
                return {'success': False, 'error': job['error']}
            
            # 检查超时
            if asyncio.get_event_loop().time() - start_time > timeout:
                return {'success': False, 'error': 'Conversion timeout'}
            
            await asyncio.sleep(0.5)
    
    def cleanup_job(self, job_id: str):
        """清理任务数据"""
        if job_id in self.jobs:
            job = self.jobs[job_id]
            # 清理临时文件
            if job.get('result') and job['result'].get('file_path'):
                try:
                    os.unlink(job['result']['file_path'])
                except OSError:
                    pass
            del self.jobs[job_id]